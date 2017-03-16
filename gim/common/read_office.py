#encoding: utf-8
import docx,xlrd,sys
reload(sys)
sys.setdefaultencoding( "utf-8" )
class office:
    def __init__(self):
        print ''
    def read_docx(self,path):
        content = ''
        doc = docx.Document(path)
        # t = doc.tables[0]
        # name = t.cell(0,0).text
        # print name
        for paragraph in doc.paragraphs:
            if len(paragraph.text) > 0:
                content += paragraph.text
                # content += '|'
        return content

    def read_docx2(self, path):
        # content = ''
        doc = docx.Document(path)
        # print len(doc.tables)
        # if len(doc.tables) > 1:
        #     t = doc.tables[index+1]
        return doc
    def read_docx2parag(self, path):
        content = ''
        doc = docx.Document(path)
        t = doc.paragraphs
        return t
    #Read excel
    def read_excel(self,path):
        xlsList = []
        data = xlrd.open_workbook(path, 'rb')
        sheet_num = len(data.sheets())
        for i in range(0,sheet_num):
            #Special add
            c0 = ''
            c1 = ''
            # c2 = ''
            #
            fullText = ''
            sheet = data.sheet_by_index(i)
            for row in range(1,sheet.nrows):
                rowData = ''
                for col in range(sheet.ncols):
                    cv = sheet.cell_value(row,col)
                    #Special add
                    if col == 0:
                        if len(cv) == 0:
                            cv = c0
                        else:
                            c0 = cv
                    if col == 1:
                        if len(cv) == 0:
                            cv = c1
                        else:
                            c1 = cv
                    #
                    rowData += str(cv) + ','
                rowData = rowData[:len(rowData)-1]
                fullText += rowData +'\r\n'
            # print fullText
            xlsList.append(fullText)
        return xlsList